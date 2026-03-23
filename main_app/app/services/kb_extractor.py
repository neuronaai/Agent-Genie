"""Knowledge-base content extraction utilities.

Handles real content extraction for URL, PDF, DOCX, and legacy DOC
knowledge-base items so that the extracted text can be compiled into the
live Retell prompt.

Dependencies (in requirements.txt):
    - requests          (HTTP fetching)
    - beautifulsoup4    (HTML parsing)
    - PyPDF2            (PDF text extraction)
    - python-docx       (DOCX text extraction)
    - olefile           (legacy .doc text extraction)
"""
from __future__ import annotations

import logging
import os
import re
import shutil
import subprocess
from typing import Optional

logger = logging.getLogger(__name__)

# ── Constants ──────────────────────────────────────────────────────────────
MAX_EXTRACTED_CHARS = 30_000          # Cap extracted text to avoid oversized prompts
REQUEST_TIMEOUT = 15                  # Seconds for HTTP requests
MAX_DOWNLOAD_BYTES = 10 * 1024 * 1024  # 10 MB — hard streaming cap
_STREAM_CHUNK_SIZE = 64 * 1024        # 64 KB chunks for streaming reads


# ── Public API ─────────────────────────────────────────────────────────────

def extract_from_url(url: str) -> Optional[str]:
    """Fetch a URL and return the main text content.

    Uses *streaming* download with a hard byte cap (``MAX_DOWNLOAD_BYTES``)
    enforced regardless of whether the server sends a ``Content-Length``
    header.  After downloading, the HTML is parsed with BeautifulSoup to
    extract the primary readable text.

    Returns ``None`` on failure.
    """
    try:
        import requests
        from bs4 import BeautifulSoup
    except ImportError as e:
        logger.error(f'Missing dependency for URL extraction: {e}')
        return None

    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (compatible; AgentGenieBot/1.0)',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        }
        resp = requests.get(url, headers=headers, timeout=REQUEST_TIMEOUT,
                            allow_redirects=True, stream=True)
        resp.raise_for_status()

        # ── Early reject if Content-Length is known and too large ──
        content_length = resp.headers.get('Content-Length')
        if content_length:
            try:
                if int(content_length) > MAX_DOWNLOAD_BYTES:
                    logger.warning(
                        f'URL content too large ({content_length} bytes, '
                        f'limit {MAX_DOWNLOAD_BYTES}): {url}'
                    )
                    resp.close()
                    return None
            except (ValueError, TypeError):
                pass  # Malformed header — fall through to streaming cap

        # ── Stream-read with hard byte cap ──
        chunks: list[bytes] = []
        downloaded = 0
        for chunk in resp.iter_content(chunk_size=_STREAM_CHUNK_SIZE):
            if not chunk:
                continue
            remaining = MAX_DOWNLOAD_BYTES - downloaded
            if remaining <= 0:
                logger.info(
                    f'Streaming cap reached ({MAX_DOWNLOAD_BYTES} bytes) '
                    f'while fetching {url} — truncating'
                )
                break
            if len(chunk) > remaining:
                chunks.append(chunk[:remaining])
                downloaded += remaining
                break
            chunks.append(chunk)
            downloaded += len(chunk)
        resp.close()

        content = b''.join(chunks)
        encoding = resp.encoding or 'utf-8'
        html = content.decode(encoding, errors='replace')

        soup = BeautifulSoup(html, 'html.parser')

        # Remove non-content elements
        for tag in soup(['script', 'style', 'nav', 'header', 'footer',
                         'aside', 'form', 'iframe', 'noscript', 'svg']):
            tag.decompose()

        # Try to find the main content area
        main = (soup.find('main')
                or soup.find('article')
                or soup.find('div', role='main')
                or soup.find('div', class_=re.compile(r'content|article|post', re.I))
                or soup.body
                or soup)

        text = main.get_text(separator='\n', strip=True)

        # Collapse excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = text[:MAX_EXTRACTED_CHARS]

        if len(text.strip()) < 50:
            logger.warning(f'Extracted text too short from URL: {url}')
            return None

        return text.strip()

    except Exception as e:
        logger.warning(f'Failed to fetch/parse URL {url}: {e}')
        return None


def extract_from_pdf(file_path: str) -> Optional[str]:
    """Extract text from a PDF file.

    Uses PyPDF2 to read all pages and concatenate the text.
    Returns ``None`` on failure.
    """
    if not os.path.isfile(file_path):
        logger.warning(f'PDF file not found: {file_path}')
        return None

    try:
        from PyPDF2 import PdfReader
    except ImportError:
        logger.error('PyPDF2 not installed — cannot extract PDF content')
        return None

    try:
        reader = PdfReader(file_path)
        pages_text: list[str] = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text.strip())
            # Safety cap on number of pages
            if i >= 200:
                pages_text.append('[... remaining pages truncated ...]')
                break

        text = '\n\n'.join(pages_text)
        text = text[:MAX_EXTRACTED_CHARS]

        if len(text.strip()) < 20:
            logger.warning(f'Extracted text too short from PDF: {file_path}')
            return None

        return text.strip()

    except Exception as e:
        logger.error(f'Failed to extract PDF {file_path}: {e}')
        return None


def extract_from_docx(file_path: str) -> Optional[str]:
    """Extract text from a DOCX file.

    Uses python-docx to read all paragraphs and tables.
    Returns ``None`` on failure.
    """
    if not os.path.isfile(file_path):
        logger.warning(f'DOCX file not found: {file_path}')
        return None

    try:
        from docx import Document
    except ImportError:
        logger.error('python-docx not installed — cannot extract DOCX content')
        return None

    try:
        doc = Document(file_path)
        parts: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(' | '.join(cells))

        text = '\n'.join(parts)
        text = text[:MAX_EXTRACTED_CHARS]

        if len(text.strip()) < 20:
            logger.warning(f'Extracted text too short from DOCX: {file_path}')
            return None

        return text.strip()

    except Exception as e:
        logger.error(f'Failed to extract DOCX {file_path}: {e}')
        return None


def extract_from_doc(file_path: str) -> Optional[str]:
    """Extract text from a legacy ``.doc`` (MS Word 97-2003) file.

    Strategy (in priority order):
    1. **antiword** — if the ``antiword`` binary is available on the system
       PATH, shell out to it.  This is the most reliable extractor for
       legacy ``.doc`` files and is available as a standard package on
       Debian/Ubuntu (``apt install antiword``).
    2. **olefile** — pure-Python OLE2 parser.  Extracts raw text streams
       from the compound-document container.  Less polished output than
       antiword but works without native binaries.
    3. **Plain-text fallback** — last resort; reads the file as UTF-8 and
       strips non-printable bytes.  Works for some ``.doc`` files that are
       actually RTF or plain text under the hood.

    Returns ``None`` if all strategies fail.
    """
    if not os.path.isfile(file_path):
        logger.warning(f'.doc file not found: {file_path}')
        return None

    # ── Strategy 1: antiword ──
    if shutil.which('antiword'):
        try:
            result = subprocess.run(
                ['antiword', '-w', '0', file_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and len(result.stdout.strip()) >= 20:
                text = result.stdout.strip()[:MAX_EXTRACTED_CHARS]
                logger.info(f'Extracted .doc via antiword ({len(text)} chars): {file_path}')
                return text
        except Exception as e:
            logger.warning(f'antiword failed for {file_path}: {e}')

    # ── Strategy 2: olefile ──
    try:
        import olefile
        if olefile.isOleFile(file_path):
            ole = olefile.OleFileIO(file_path)
            # The main text stream in a Word .doc is "WordDocument" but
            # the plain-text content is stored in the "1Table" or "0Table"
            # streams — however the easiest portable approach is to read
            # all streams and extract printable text.
            text_parts: list[str] = []
            for stream_path in ole.listdir():
                stream_name = '/'.join(stream_path)
                try:
                    raw = ole.openstream(stream_path).read()
                    # Attempt UTF-16-LE decode (common in .doc text streams)
                    try:
                        decoded = raw.decode('utf-16-le', errors='ignore')
                    except Exception:
                        decoded = raw.decode('utf-8', errors='ignore')
                    # Keep only printable runs of 4+ characters
                    runs = re.findall(r'[\x20-\x7E\n\r\t]{4,}', decoded)
                    if runs:
                        text_parts.append(' '.join(runs))
                except Exception:
                    continue
            ole.close()
            if text_parts:
                combined = '\n'.join(text_parts)
                # De-duplicate repeated whitespace
                combined = re.sub(r'\s{3,}', '  ', combined)
                combined = combined[:MAX_EXTRACTED_CHARS]
                if len(combined.strip()) >= 20:
                    logger.info(f'Extracted .doc via olefile ({len(combined)} chars): {file_path}')
                    return combined.strip()
    except ImportError:
        logger.info('olefile not installed — skipping OLE extraction for .doc')
    except Exception as e:
        logger.warning(f'olefile extraction failed for {file_path}: {e}')

    # ── Strategy 3: plain-text fallback ──
    logger.warning(f'Falling back to plain-text read for .doc: {file_path}')
    return _extract_printable_text(file_path)


def _extract_printable_text(file_path: str) -> Optional[str]:
    """Read a file in binary mode and extract printable ASCII/UTF-8 runs.

    This is a last-resort extractor for binary files that may contain
    embedded text (e.g., some ``.doc`` files that are really RTF).
    """
    try:
        with open(file_path, 'rb') as f:
            raw = f.read(MAX_DOWNLOAD_BYTES)
        # Try UTF-8 first
        decoded = raw.decode('utf-8', errors='ignore')
        # Keep printable runs of 4+ characters
        runs = re.findall(r'[\x20-\x7E\n\r\t]{4,}', decoded)
        text = ' '.join(runs)
        text = re.sub(r'\s{3,}', '  ', text)
        text = text[:MAX_EXTRACTED_CHARS]
        if len(text.strip()) >= 20:
            return text.strip()
        return None
    except Exception as e:
        logger.error(f'Plain-text fallback failed for {file_path}: {e}')
        return None


def extract_from_text_file(file_path: str) -> Optional[str]:
    """Extract text from a plain text file (TXT, CSV, MD, etc.).

    Returns ``None`` on failure.
    """
    if not os.path.isfile(file_path):
        logger.warning(f'Text file not found: {file_path}')
        return None

    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read(MAX_EXTRACTED_CHARS)
        if len(text.strip()) < 20:
            return None
        return text.strip()
    except Exception as e:
        logger.error(f'Failed to read text file {file_path}: {e}')
        return None


def extract_content(*, url: str = '', file_path: str = '',
                    file_name: str = '', file_mime: str = '') -> Optional[str]:
    """Unified extraction entry point.

    Determines the best extraction strategy based on the available inputs:
    - If a URL is provided, fetches and extracts webpage content.
    - If a file is provided, dispatches to the appropriate extractor
      based on file extension or MIME type.

    Returns the extracted text, or ``None`` if extraction fails.
    """
    # URL extraction
    if url:
        return extract_from_url(url)

    # File extraction
    if file_path and os.path.isfile(file_path):
        ext = os.path.splitext(file_name or file_path)[1].lower()
        mime = (file_mime or '').lower()

        if ext == '.pdf' or 'pdf' in mime:
            return extract_from_pdf(file_path)
        elif ext in ('.docx',) or 'wordprocessingml' in mime:
            return extract_from_docx(file_path)
        elif ext == '.doc' or 'msword' in mime:
            return extract_from_doc(file_path)
        elif ext in ('.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm'):
            return extract_from_text_file(file_path)
        else:
            # Unknown type — try plain text as last resort
            logger.info(f'Unknown file type ({ext}) — attempting plain text read: {file_path}')
            return extract_from_text_file(file_path)

    return None
